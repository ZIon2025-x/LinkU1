"""生成 Link2Ur 达人入驻介绍 Word 文档(商务完善版)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.styles['Normal'].font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cn(run, size=11, bold=False, color=None):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def title(text, size=22, color=RGBColor(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cn(run, size=size, bold=True, color=color)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn(run, size=13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn(run, size=11)


def bullet(text, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    if prefix:
        r1 = p.add_run(prefix)
        set_cn(r1, size=11, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
        r2 = p.add_run(text)
        set_cn(r2, size=11)
    else:
        r = p.add_run(text)
        set_cn(r, size=11)


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_cn(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E75B6')
        cell._tc.get_or_add_tcPr().append(shd)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_cn(run, size=10)
    if col_widths:
        for row in t.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)


# ============ 封面 ============
doc.add_paragraph()
doc.add_paragraph()
title("Link2Ur  技能互助平台", size=26)
doc.add_paragraph()
title("达人入驻合作说明", size=18, color=RGBColor(0x2E, 0x75, 0xB6))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("—— 致有意入驻的个人达人与商家合作伙伴 ——")
set_cn(run, size=12, color=RGBColor(0x80, 0x80, 0x80))
doc.add_page_break()

# ============ 一、平台介绍 ============
h1("一、平台介绍")
para(
    "Link2Ur(技能互助平台)是一款面向海外华人社区的综合性本地生活服务平台。"
    "我们的目标是把身边有技能、有资源的华人与有需求的用户高效地连接起来,"
    "让一次家教、一次搬家、一次代购、一次剪发、一次咨询、一次同城聚会,"
    "都能在一个 App 内完成:从浏览、沟通、下单、支付到售后评价,全流程透明可追溯。"
)
para(
    "平台覆盖 iOS 与 Android,支持简体中文、繁体中文、英文三语切换,"
    "所有交易资金通过合规支付通道清算,直接结算到达人本人或商家团队账户。"
)

h2("平台定位")
bullet("面向海外华人(留学生、工作族、新移民、华人家庭)的一站式本地生活服务入口。")
bullet("融合「任务互助 + 达人服务 + 活动社群 + 二手交易 + 内容社区」五大板块。")
bullet("强调华人圈内的「熟人感与可信感」——共同语言、共同文化、同胞互助。")

# ============ 二、平台核心功能 ============
h1("二、平台核心功能")
para("用户在一个 App 内即可完成生活中几乎所有日常需求,主要板块如下:")

h2("1. 任务互助")
bullet("用户发布任务(搬家、代取、接送、陪同、跑腿、翻译等),达人抢单或被智能匹配。")
bullet("支持三种定价模式:固定价、可议价、待报价(由达人出价)。")
bullet("支持灵活任务与限时任务、线上 / 线下 / 混合三种执行方式。")
bullet("支持多人同时参与的「多人任务」(适合大扫除、活动布置等)。")
bullet("完整的议价 / 反报价 / 申请审批流程,杜绝口头约定扯皮。")
bullet("5 天自动确认机制、完工提醒、纠纷仲裁与退款保障。")

h2("2. 达人服务市场(本项目核心)")
bullet("达人可将自己的技能打包成标准化「服务套餐」,明码标价,用户像逛淘宝一样下单。")
bullet("服务支持图文展示、作品案例、技能标签、分类(教育、翻译、美业、餐饮、编程、手工、娱乐、代办等)。")
bullet("支持时段预约——达人设置可接单的日期/时段,用户按日历下单,系统自动提醒。")
bullet("支持服务咨询(Consultation)——用户先询价沟通,达人发送正式报价后再下单。")
bullet("支持官方认证、达人等级、达人徽章,帮助用户一眼识别靠谱服务者。")

h2("3. 活动与拼单")
bullet("达人可开设多人参与的活动(课程、workshop、聚餐、户外、观影等),设定人数上限与截止时间。")
bullet("支持拼单成团:凑够最少人数自动成团,享折扣价,适合课程包与团购场景。")
bullet("支持抽奖活动:奖品可为积分、实物、兑换码、线下参与资格等。")
bullet("活动支持现金 / 积分 / 混合付费,奖励可均分或按贡献分配。")
bullet("活动可按固定日期或循环时段发布,方便长期运营。")

h2("4. 跳蚤市场")
bullet("面向留学生与华人社区的本地二手交易、闲置转让、以物换物。")
bullet("支持议价、反报价、求购请求、租赁模式与物品咨询。")
bullet("举报与内容审核机制,保障交易环境安全。")

h2("5. 内容社区与论坛")
bullet("分版块论坛(生活、学习、求助、兴趣、吐槽等),支持发帖、跟帖、嵌套回复、点赞、收藏。")
bullet("达人可通过发帖、回答问题、分享作品获得自然曝光,内容可直接挂链到店铺或服务。")
bullet("完整的举报、置顶、精华、锁帖等社区治理能力。")

h2("6. 即时通讯 & AI 助手")
bullet("任务聊天、私信、客服会话三类 IM 通道,支持图片、文件、系统消息、议价消息。")
bullet("AI 智能助手帮用户自动生成任务草稿、智能匹配达人,降低新用户下单门槛。")
bullet("智能推荐引擎根据兴趣、技能、位置为用户推送合适的达人与任务。")

h2("7. 钱包、优惠券与积分")
bullet("每位用户拥有独立钱包,查看余额、收益、消费、退款全记录。")
bullet("积分体系:每日签到、任务完成、活动参与均可累积积分,兑换优惠券或商品。")
bullet("完整的优惠券体系:满减、折扣、限领、限用、叠加规则、VAT 分类一应俱全。")

h2("8. 社交与成长体系")
bullet("关注 / 粉丝、主页展示、浏览记录、技能认可、徽章成就。")
bullet("技能排行榜、达人榜单、用户等级(普通 / VIP / 超级),激励长期运营。")
bullet("新手任务、签到连击奖励、新人礼包等 Gamification 玩法提升活跃。")

h2("9. 学生认证")
bullet("通过学校邮箱验证学生身份,获得「已认证学生」徽章,享受学生专属优惠与信任加成。")
bullet("支持多所海外高校,支持续期与邮箱变更。")

# ============ 三、达人功能详解 ============
doc.add_page_break()
h1("三、达人功能详解(核心重点)")
para(
    "Link2Ur 在 2026 年完成了「达人体系」的全面升级,投入了平台最大规模的产品与运营资源。"
    "不同于普通信息撮合平台,Link2Ur 为达人提供的是一整套「线上店铺 + 经营后台 + 营销工具」"
    "的完整运营解决方案。"
)

h2("1. 个人达人 vs 商家团队 双模式")
bullet(
    "适合留学生、自由职业者、手艺人、斜杠青年。实名认证后即可入驻,独立接单、独立钱包、独立店铺。",
    "个人达人:")
bullet(
    "适合工作室、门店、机构、品牌方。支持多成员协作,角色分为店主 / 管理员 / 员工,"
    "统一品牌展示,订单可按成员分派,团队共享钱包与结算账户。", "商家团队(达人团队):")
bullet("一名用户可同时是个人达人,也可加入/创建多个商家团队,灵活兼顾个人与组织身份。", "灵活身份:")

h2("2. 服务商品化能力")
bullet("将技能打包为明码标价的服务套餐,含图文、案例、时长、地点(线上/线下/均可)。")
bullet("支持服务分类、技能标签、服务地点类型、多语言描述(自动翻译)。")
bullet("支持服务包(Package)——多个服务捆绑销售,提升客单价。")
bullet("支持定价灵活性:固定价 / 可议价 / 待报价。")

h2("3. 排班与预约系统")
bullet("达人可设置可接单时段(按日、按小时),支持循环排班与单次日程。")
bullet("支持设置休息日 / 不接单日 / 节假日黑名单,避免打扰。")
bullet("用户按日历选时段下单,系统自动确认、自动提醒双方。")

h2("4. 咨询与议价流程")
bullet("用户对服务感兴趣时,可先发起咨询(Consultation),向达人询价沟通。")
bullet("达人可发送正式报价 / 反报价,双方达成一致后用户下单付款。")
bullet("整个流程有完整记录,避免「聊完就忘」「口头承诺不算数」的问题。")

h2("5. 活动与拼单运营")
bullet("达人可发起多人活动或拼单团购,凑够最少人数即成团。")
bullet("活动可附带抽奖、积分奖励、老带新裂变机制。")
bullet("支持限时 / 限量 / 倒计时,适合做节日促销与社群运营。")

h2("6. 达人专属店铺页")
bullet("每位达人 / 每个商家团队拥有独立店铺主页,展示头像、简介、作品、服务列表、评价墙、粉丝数。")
bullet("用户可「关注达人」,达人上新、促销、开团时第一时间推送通知。")
bullet("商家团队页面支持成员展示与角色标注,塑造专业品牌形象。")

h2("7. 经营后台(Dashboard)")
bullet("实时查看订单、待接单、进行中、已完成、收入明细、转化率、好评率等核心指标。")
bullet("完整的订单生命周期管理:接单 → 沟通 → 进度 → 完工 → 评价 → 售后。")
bullet("支持商家回复评价,主动塑造口碑(避免差评无处申辩)。")
bullet("提交「资料更新申请」修改认证信息,平台审核后生效,保障资质严肃性。")

h2("8. 营销与曝光工具")
bullet("首页推荐位、分类置顶、搜索加权、发现页推送多通道曝光。")
bullet("自主发放优惠券、满减券、新客券、粉丝专属券。")
bullet("参与平台活动(节日大促、新人扶持、周末特惠)获得额外流量倾斜。")
bullet("在论坛、活动、动态中发帖种草,直接导流到店铺与服务页。")

h2("9. 收入与提现")
bullet("交易款项通过合规支付通道清算,资金透明可追溯,规避微信转账的合规风险。")
bullet("达人钱包余额可随时申请提现,1 - 3 个工作日到账。")
bullet("支持英镑等多币种结算(未来扩展更多地区)。")
bullet("每笔交易均生成电子账单,支持月度对账与发票申请,方便报税与做账。")

h2("10. 评价与信任体系")
bullet("任务 / 服务 / 活动完成后,用户可打星评价 + 文字评价。")
bullet("达人可对评价做专业回复,树立服务形象。")
bullet("系统综合好评率、完成率、响应速度计算达人评分,作为推荐权重。")
bullet("平台提供官方认证徽章、学生认证徽章、商家认证徽章,提升用户信任度。")

# ============ 四、目标客户与市场 ============
doc.add_page_break()
h1("四、目标客户与市场分析")

h2("1. 需求方用户画像(谁在买)")
bullet("在英国 / 欧洲 / 北美的中国留学生(18 - 28 岁),高频需求:搬家、代取、家教、代购、接送机、作业辅导。")
bullet("海外华人家庭与新移民,侧重家庭场景:家政、月嫂、装修、子女中文教育、税务法律咨询。")
bullet("海外工作的华人白领,侧重个人提升与生活品质:健身、美业、心理咨询、同城社交、兴趣班。")
bullet("本地华人社群 / 学生会 / 兴趣组织,需要活动组织、场地、物料、宣传等对接服务。")

h2("2. 供给方达人画像(谁在卖)")
bullet("有一技之长、希望利用业余时间变现的留学生与斜杠青年。")
bullet("本地华人工作室与小型门店:美甲美睫、餐饮烘焙、教培机构、家政清洁、维修搬家、摄影跟拍。")
bullet("专业服务从业者:翻译、税务、法律、心理、医疗陪同、留学咨询、签证顾问。")
bullet("希望拓展华人客群的本地非华人商家(平台原生多语言,非华人背景也可入驻)。")

h2("3. 市场机会与规模")
para(
    "仅英国一地,在校中国留学生长期维持在 15 万人以上;加上工作签、陪读签、新移民、"
    "华人家庭与二代华人,海外华人规模庞大且持续增长。目标人群具有三个显著特征:"
)
bullet("语言与文化壁垒 —— 更倾向在华人圈内寻找服务,熟人信任与沟通成本远优于本地平台。")
bullet("需求高频且分散 —— 日常生活、求学、求职、社交的需求被大量微信群 / 朋友圈割裂,缺少统一入口。")
bullet("付费意愿强 —— 相比国内同龄人,海外华人对「省时、省心、靠谱」的服务支付意愿更高。")
para(
    "目前海外华人市场虽然存在一些信息群、朋友圈代购、同城小程序,但普遍「无保障、"
    "无闭环、无售后」,纠纷频发、体验割裂。Link2Ur 的定位正是填补这一空白——"
    "让华人之间的服务交易像淘宝 / 美团一样标准、安全、可评价、可追溯。"
)

h2("4. 我们的差异化优势")
bullet("专注华人社区,用户画像精准,获客成本显著低于投放公域流量。")
bullet("服务 + 任务 + 活动 + 二手 + 社区五合一,单用户 LTV 高、场景覆盖完整。")
bullet("原生支持商家团队模式,而不只是个人兼职工具——商家可直接把线下门店搬上线。")
bullet("合规资金通道 + 完整售后 + 争议仲裁,彻底摆脱微信私转带来的纠纷与合规风险。")

# ============ 五、合作方式与入驻流程 ============
doc.add_page_break()
h1("五、合作方式与入驻流程")

h2("1. 入驻身份类型")
bullet("适合有一技之长的个人、留学生、自由职业者。实名认证即可入驻,无需营业执照。", "个人达人:")
bullet(
    "适合工作室、门店、机构、品牌方。可多人协作、品牌化展示、团队分账、统一对账。", "商家团队:")
bullet(
    "面向社群组织者、校园大使、KOL、本地华人公众号 / 小红书达人,与平台共建内容生态。",
    "内容/战略合作伙伴:")

h2("2. 入驻四步流程")
bullet("下载 Link2Ur App,使用手机号 / 邮箱注册账号,完成基础实名认证。", "第一步 注册:")
bullet(
    "在「我的 - 成为达人 / 商家入驻」提交资料(身份信息、服务类目、作品案例、资质证书)。",
    "第二步 申请:")
bullet(
    "平台运营团队审核(通常 1 - 3 个工作日),通过后自动开通达人后台与店铺页。",
    "第三步 审核:")
bullet(
    "上架服务套餐、配置时段、完善店铺简介与作品,发布首个服务即开始接单运营。",
    "第四步 上线:")

h2("3. 达人核心权益")
bullet("零入驻费、零年费、零保证金(特殊高风险品类除外)。")
bullet("完整的店铺页、经营后台、营销工具全免费使用。")
bullet("新人达人享首 30 天「流量扶持期」——首页曝光、新客券补贴、平台代发推广。")
bullet("专属运营对接一对一指导首单落地。")
bullet("优秀达人纳入「平台精选」和「达人推荐榜」,享受长期流量倾斜。")

# ============ 六、抽成方案与结算政策 ============
h1("六、抽成方案与结算政策")

para(
    "Link2Ur 采用「零入驻费 + 按成交抽佣」的轻模式,达人零前期成本,"
    "仅在订单成功完成后由平台收取服务佣金,与达人利益完全一致。"
)

h2("1. 平台抽佣比例")
table(
    headers=["业务类别", "平台抽佣", "说明"],
    rows=[
        ["达人服务(接单、套餐、预约、咨询)", "8%", "包括个人达人与商家团队的所有服务订单"],
        ["达人活动(活动报名、拼单、课程团购)", "6%", "鼓励达人开展多人活动与社群运营"],
        ["二手交易(跳蚤市场)", "0%", "纯 C2C 撮合,仅收支付通道手续费"],
        ["任务互助(通用任务)", "另行约定", "普通用户间的互助任务,按单独规则计算"],
    ],
    col_widths=[7.5, 2.5, 6.5],
)
doc.add_paragraph()
para(
    "说明:以上抽佣比例按订单实际成交金额(用户实际支付金额扣除优惠券部分后)计算,"
    "透明公开、不设隐藏收费。所有订单账单与佣金明细在达人后台可实时查询。"
)

h2("2. 支付通道手续费")
bullet("所有线上支付另由第三方合规支付通道收取约 1.5% - 2.9% + 固定小额手续费,由平台代扣代缴,账单可查。")
bullet("此部分为支付机构收取,平台不从中获利。")

h2("3. 新人扶持政策")
bullet("新达人入驻后前 30 天内成交的订单,平台抽佣再减免 50%(服务 4% / 活动 3%)。")
bullet("新达人前 10 单自动进入首页新达人推荐位,获得启动流量。")
bullet("商家团队前 3 个月享受「0 佣金体验期」(仅付支付通道手续费),帮助快速起量。")
bullet("老达人推荐新达人入驻,双方各得奖励积分与优惠券。")

h2("4. 结算与提现")
bullet("订单完成并经用户确认(或自动确认期满)后,款项进入达人钱包。")
bullet("达人可随时申请提现,常规到账周期 1 - 3 个工作日。")
bullet("每笔交易均生成电子账单,支持月度对账、导出流水与发票申请,方便达人记账报税。")

h2("5. 可选增值服务")
bullet("首页 Banner、分类置顶、开屏推荐等付费广告位,按周 / 月购买,提升曝光。", "付费推广:")
bullet("平台代运营、内容代写、活动策划,适合希望省心经营的商家团队。", "代运营服务:")
bullet("参与平台双 11、新春、开学季等大促,享受专属流量与补贴。", "联合营销:")

para("(以上政策为当前标准方案,正式合作以双方签署的《达人合作协议》为准,平台保留根据市场情况调整的权利。)")

# ============ 七、为什么选择 Link2Ur ============
doc.add_page_break()
h1("七、为什么选择 Link2Ur")
bullet("明确的海外华人社区定位 —— 用户精准、获客成本低、天然信任。")
bullet("一站式交易闭环 —— 曝光、沟通、下单、支付、售后一条龙,告别信息群的割裂体验。")
bullet("原生商家团队模式 —— 不只是兼职工具,而是可承载整个线下门店的数字化经营平台。")
bullet("低佣金率 —— 达人服务 8%、活动 6%,显著低于主流本地生活平台。")
bullet("合规资金通道 —— 彻底摆脱微信转账的合规与纠纷风险。")
bullet("零入驻成本 + 新人扶持 + 商家免佣期 —— 试水零风险,起量有扶持。")
bullet("完整的经营工具链 —— 店铺、套餐、排班、营销、数据、评价管理一应俱全。")
bullet("与平台共同成长 —— 早期入驻的达人享受长期流量红利与合作优先权。")

# ============ 八、联系我们 ============
h1("八、联系我们")
para("如您希望进一步了解合作细节或预约入驻洽谈,欢迎通过以下方式联系我们:")
bullet("官方邮箱:  [待填写]")
bullet("商务合作微信 / 电话:  [待填写]")
bullet("官方网站:  [待填写]")
bullet("App 下载:  应用商店搜索「Link2Ur」")
bullet("客服支持:  App 内「我的 - 联系客服」")

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("—— 让每一份技能都被看见,让每一次服务都有回响 ——")
set_cn(run, size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Link2Ur 期待与您同行")
set_cn(run, size=14, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))

out = r"F:\python_work\LinkU\docs\Link2Ur_达人入驻介绍.docx"
doc.save(out)
print(f"SAVED: {out}")
